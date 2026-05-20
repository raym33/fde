"""Extracción real de documentos para RAG.

Soporta TXT/MD, DOCX y PDF. Para PDFs escaneados intenta OCR opcional con
PyMuPDF + pytesseract si Tesseract está instalado en el sistema.
"""
from __future__ import annotations

import io
import shutil
from dataclasses import dataclass, field
from pathlib import Path


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}
TEXT_CONTENT_TYPES = {"text/plain", "text/markdown", "application/octet-stream"}
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
PDF_CONTENT_TYPES = {"application/pdf"}


class DocumentParseError(ValueError):
    """Documento no soportado o imposible de parsear."""


@dataclass
class ExtractedDocument:
    text: str
    parser: str
    metadata: dict = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)


def parser_status() -> dict:
    return {
        "supported_extensions": sorted(SUPPORTED_EXTENSIONS),
        "pdf_text": _module_available("pypdf"),
        "docx": _module_available("docx"),
        "ocr": {
            "pymupdf": _module_available("fitz"),
            "pytesseract": _module_available("pytesseract"),
            "tesseract_binary": bool(shutil.which("tesseract")),
            "available": _ocr_available(),
        },
    }


def parse_document(
    raw: bytes,
    *,
    filename: str,
    content_type: str | None = None,
    ocr_if_needed: bool = True,
) -> ExtractedDocument:
    ext = Path(filename).suffix.lower()
    content_type = (content_type or "").split(";")[0].lower()

    # Prioriza extensión fiable: muchos clientes/subidas mandan DOCX/PDF como
    # application/octet-stream.
    if ext == ".docx":
        return _parse_docx(raw, filename=filename)
    if ext == ".pdf":
        return _parse_pdf(raw, filename=filename, ocr_if_needed=ocr_if_needed)
    if ext in {".txt", ".md", ".markdown"}:
        return _parse_text(raw, filename=filename)
    if content_type in DOCX_CONTENT_TYPES:
        return _parse_docx(raw, filename=filename)
    if content_type in PDF_CONTENT_TYPES:
        return _parse_pdf(raw, filename=filename, ocr_if_needed=ocr_if_needed)
    if content_type in TEXT_CONTENT_TYPES:
        return _parse_text(raw, filename=filename)
    raise DocumentParseError(
        f"Formato no soportado: {filename or content_type}. Soportados: TXT, MD, DOCX, PDF."
    )


def _parse_text(raw: bytes, *, filename: str) -> ExtractedDocument:
    warnings = []
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = raw.decode(encoding)
            return ExtractedDocument(
                text=_clean_text(text),
                parser=f"text/{encoding}",
                metadata={"filename": filename, "format": "text", "bytes": len(raw)},
                warnings=warnings,
            )
        except UnicodeDecodeError:
            continue
    text = raw.decode("latin-1", errors="ignore")
    warnings.append("No se pudo detectar encoding con precisión; se usó latin-1 con ignore.")
    return ExtractedDocument(
        text=_clean_text(text),
        parser="text/latin-1-fallback",
        metadata={"filename": filename, "format": "text", "bytes": len(raw)},
        warnings=warnings,
    )


def _parse_docx(raw: bytes, *, filename: str) -> ExtractedDocument:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        raise DocumentParseError("python-docx no está instalado.") from exc

    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = _clean_text("\n\n".join(parts))
    if not text:
        raise DocumentParseError("DOCX sin texto extraíble.")
    return ExtractedDocument(
        text=text,
        parser="python-docx",
        metadata={
            "filename": filename,
            "format": "docx",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "bytes": len(raw),
        },
    )


def _parse_pdf(raw: bytes, *, filename: str, ocr_if_needed: bool) -> ExtractedDocument:
    warnings = []
    text_parts: list[str] = []
    pages = 0
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        pages = len(reader.pages)
        for index, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"\n\n[page {index}]\n{page_text}")
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"Extracción PDF textual falló: {exc}")

    text = _clean_text("\n".join(text_parts))
    if len(text) >= 80 or not ocr_if_needed:
        if not text:
            raise DocumentParseError("PDF sin texto extraíble.")
        return ExtractedDocument(
            text=text,
            parser="pypdf",
            metadata={"filename": filename, "format": "pdf", "pages": pages, "bytes": len(raw)},
            warnings=warnings,
        )

    if not _ocr_available():
        warnings.append("PDF con poco texto; OCR no disponible (instala tesseract).")
        if text:
            return ExtractedDocument(
                text=text,
                parser="pypdf-partial",
                metadata={"filename": filename, "format": "pdf", "pages": pages, "bytes": len(raw)},
                warnings=warnings,
            )
        raise DocumentParseError("PDF escaneado o sin texto; OCR no disponible.")

    ocr_text, ocr_pages = _ocr_pdf(raw)
    text = _clean_text(ocr_text)
    if not text:
        raise DocumentParseError("OCR ejecutado, pero no se extrajo texto.")
    warnings.append("Texto extraído mediante OCR; revisar precisión en documentos críticos.")
    return ExtractedDocument(
        text=text,
        parser="pymupdf+pytesseract",
        metadata={
            "filename": filename,
            "format": "pdf",
            "pages": ocr_pages or pages,
            "bytes": len(raw),
            "ocr": True,
        },
        warnings=warnings,
    )


def _ocr_pdf(raw: bytes) -> tuple[str, int]:
    import fitz
    import pytesseract
    from PIL import Image

    doc = fitz.open(stream=raw, filetype="pdf")
    parts = []
    for index, page in enumerate(doc, start=1):
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(image, lang="spa+eng")
        if text.strip():
            parts.append(f"\n\n[page {index} OCR]\n{text}")
    return "\n".join(parts), len(doc)


def _ocr_available() -> bool:
    return (
        _module_available("fitz")
        and _module_available("pytesseract")
        and _module_available("PIL")
        and bool(shutil.which("tesseract"))
    )


def _module_available(name: str) -> bool:
    try:
        __import__(name)
        return True
    except Exception:
        return False


def _clean_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.replace("\x00", "").splitlines()]
    return "\n".join(line for line in lines if line).strip()
